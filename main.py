# importing libraries
import re
import requests
import bs4
import docx
import pandas as pd
import oxford as ox


# importing text
doc = docx.Document('test.docx')
a = doc.paragraphs[0].text


# extracting words
pattern = '([A-Za-z]{4,})'
w_dict = re.findall(pattern, a)


# formatting
c = 0
r_dict = {}
for i in w_dict:
    if len(r_dict) < 10:
        try:
            ox.Word.get(i)
        except:
            continue
        wordform = ox.Word.wordform()
        try:
            definition = ox.Word.definitions()[0]
            r_dict.update({i: [wordform, definition]})
        except:
            continue
    else:
        break
pd_dict = pd.DataFrame.from_dict(r_dict, orient='index')


# exporting to word

out = docx.Document()
out.add_heading('Glossary', 0)
for i in r_dict:
    a = out.add_paragraph()
    a.add_run(i.capitalize()).font.bold = True
    a.add_run(' - ')
    a.add_run(r_dict[i][0]).font.italic = True
    a.add_run(' - ')
    a.add_run(r_dict[i][1])
out.save('glossary.docx')

print([i.text for i in out.paragraphs])